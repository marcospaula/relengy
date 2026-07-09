"""Relatório de análise: descreve-se uma vez, sai em Markdown e em Word.

Um estudo de confiabilidade precisa virar documento — alguém assina, alguém audita,
alguém lê daqui a três anos. Escrever o mesmo texto duas vezes (uma no notebook,
outra no Word) é como manter duas planilhas de FMEA: elas divergem.

Aqui o relatório é uma lista de blocos. `to_markdown` renderiza para o GitHub;
`to_docx` para quem quer o arquivo no e-mail. Os dois leem a mesma fonte.

    r = Report("Bomba centrífuga P-101", subtitle="FMEA x 5 Porquês")
    r.h1("Achados").para("...").table(df).image("fta.png", "Árvore de falhas")
    r.to_markdown("RELATORIO.md"); r.to_docx("RELATORIO.docx")

Imagens em Markdown ficam relativas ao diretório do .md; no .docx são embutidas.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

import pandas as pd

BlockKind = Literal["h1", "h2", "h3", "para", "bullets", "table", "image", "pre", "callout"]


@dataclass
class Block:
    kind: BlockKind
    payload: object
    caption: str | None = None


@dataclass
class Report:
    title: str
    subtitle: str | None = None
    base_dir: Path = field(default_factory=Path.cwd)
    """Diretório contra o qual caminhos de imagem relativos são resolvidos.

    O Markdown reescreve o caminho relativo ao .md; o Word embute o arquivo e
    precisa dele resolvido. Guardar a base evita que os dois divirjam.
    """
    blocks: list[Block] = field(default_factory=list)

    # -- construção (encadeável) -------------------------------------------

    def _add(self, kind: BlockKind, payload: object, caption: str | None = None) -> Report:
        self.blocks.append(Block(kind, payload, caption))
        return self

    def h1(self, t: str) -> Report:
        return self._add("h1", t)

    def h2(self, t: str) -> Report:
        return self._add("h2", t)

    def h3(self, t: str) -> Report:
        return self._add("h3", t)

    def para(self, t: str) -> Report:
        return self._add("para", t)

    def bullets(self, items: list[str]) -> Report:
        return self._add("bullets", list(items))

    def pre(self, text: str) -> Report:
        return self._add("pre", text)

    def callout(self, text: str) -> Report:
        """Bloco de destaque: citação no Markdown, itálico recuado no Word."""
        return self._add("callout", text)

    def table(self, df: pd.DataFrame, caption: str | None = None, floatfmt: str = "%.4g") -> Report:
        return self._add("table", (df.copy(), floatfmt), caption)

    def image(self, path: str | Path, caption: str | None = None, width_in: float = 6.0) -> Report:
        p = Path(path)
        if not p.is_absolute():
            p = self.base_dir / p
        return self._add("image", (p.resolve(), width_in), caption)

    # -- saída --------------------------------------------------------------

    def to_markdown(self, path: str | Path) -> Path:
        path = Path(path)
        base = path.parent
        out = [f"# {self.title}", ""]
        if self.subtitle:
            out += [f"*{self.subtitle}*", ""]

        for b in self.blocks:
            if b.kind == "h1":
                out += [f"## {b.payload}", ""]
            elif b.kind == "h2":
                out += [f"### {b.payload}", ""]
            elif b.kind == "h3":
                out += [f"#### {b.payload}", ""]
            elif b.kind == "para":
                out += [str(b.payload), ""]
            elif b.kind == "callout":
                out += [f"> {str(b.payload)}", ""]
            elif b.kind == "bullets":
                out += [f"- {i}" for i in b.payload] + [""]
            elif b.kind == "pre":
                out += ["```text", str(b.payload).rstrip(), "```", ""]
            elif b.kind == "table":
                df, fmt = b.payload
                if b.caption:
                    out += [f"**{b.caption}**", ""]
                out += [_md_table(df, fmt), ""]
            elif b.kind == "image":
                p, _ = b.payload
                rel = os.path.relpath(p, base)
                out += [f"![{b.caption or ''}]({rel})", ""]
                if b.caption:
                    out += [f"*{b.caption}*", ""]

        path.write_text("\n".join(out), encoding="utf-8")
        return path

    def to_docx(self, path: str | Path) -> Path:
        from docx import Document
        from docx.shared import Inches, Pt

        path = Path(path)
        doc = Document()
        doc.add_heading(self.title, level=0)
        if self.subtitle:
            p = doc.add_paragraph(self.subtitle)
            p.runs[0].italic = True

        for b in self.blocks:
            if b.kind in ("h1", "h2", "h3"):
                doc.add_heading(str(b.payload), level={"h1": 1, "h2": 2, "h3": 3}[b.kind])
            elif b.kind == "para":
                doc.add_paragraph(str(b.payload))
            elif b.kind == "callout":
                p = doc.add_paragraph(str(b.payload), style="Intense Quote")
                del p  # o estilo já diz tudo
            elif b.kind == "bullets":
                for i in b.payload:
                    doc.add_paragraph(str(i), style="List Bullet")
            elif b.kind == "pre":
                p = doc.add_paragraph()
                run = p.add_run(str(b.payload).rstrip())
                run.font.name = "Consolas"
                run.font.size = Pt(8)
            elif b.kind == "table":
                df, fmt = b.payload
                if b.caption:
                    cp = doc.add_paragraph(b.caption)
                    cp.runs[0].bold = True
                _docx_table(doc, df, fmt)
            elif b.kind == "image":
                p, width_in = b.payload
                doc.add_picture(str(p), width=Inches(width_in))
                if b.caption:
                    cap = doc.add_paragraph(b.caption)
                    cap.runs[0].italic = True

        doc.save(path)
        return path


# -- helpers ----------------------------------------------------------------

def _fmt(v: object, floatfmt: str) -> str:
    try:
        if v is None or pd.isna(v):
            return "—"
    except (TypeError, ValueError):
        pass  # arrays e afins: cai no str()
    if isinstance(v, float):
        return floatfmt % v
    return str(v)


def _md_table(df: pd.DataFrame, floatfmt: str) -> str:
    cols = [str(c) for c in df.columns]
    head = "| " + " | ".join(cols) + " |"
    sep = "|" + "|".join("---" for _ in cols) + "|"
    rows = [
        "| " + " | ".join(_fmt(v, floatfmt).replace("|", "\\|") for v in r) + " |"
        for r in df.itertuples(index=False)
    ]
    return "\n".join([head, sep, *rows])


def _docx_table(doc, df: pd.DataFrame, floatfmt: str) -> None:
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for c, name in zip(t.rows[0].cells, df.columns):
        c.text = str(name)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
    for r in df.itertuples(index=False):
        cells = t.add_row().cells
        for c, v in zip(cells, r):
            c.text = _fmt(v, floatfmt)
